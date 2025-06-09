# This file has been refactored into smaller modules.
# Import the new modular ResumeProcessor for backward compatibility
from processors.resume_processor import ResumeProcessor

# For backward compatibility, expose the ResumeProcessor class
__all__ = ['ResumeProcessor']

import os
import re
import json
import hashlib
import logging
from pathlib import Path
from datetime import datetime, timedelta
from openai import OpenAI
from typing import List, Dict, Optional, Any
import PyPDF2
from docx import Document
from config_loader import get_config
import pandas as pd
from concurrent.futures import ThreadPoolExecutor, as_completed
import time

class ResumeProcessor:
    def __init__(self, cache_dir: str = None):
        self.config = get_config()
        self.client = OpenAI(api_key=self.config.get_openai_api_key())
        self.cache_dir = Path(cache_dir or self.config.get_cache_directory())
        
        # Set up logging for this class
        self.logger = logging.getLogger(f"{__name__}.{self.__class__.__name__}")
        
        self._ensure_cache_directory()
        
    def _ensure_cache_directory(self):
        """Ensure the cache directory exists"""
        try:
            self.cache_dir.mkdir(exist_ok=True)
            self.logger.debug(f"Cache directory ensured: {self.cache_dir}")
        except Exception as e:
            self.logger.warning(f"Could not create cache directory {self.cache_dir}: {e}")
            # Fall back to current directory if cache creation fails
            self.cache_dir = Path(".")
            self.logger.info(f"Falling back to current directory for cache: {self.cache_dir}")

    def _generate_cache_key(self, content: str, operation: str, **kwargs) -> str:
        """Generate a unique cache key based on content and parameters"""
        # Create a string that includes content + operation + any additional parameters
        cache_input = f"{operation}:{content}"
        for key, value in sorted(kwargs.items()):
            cache_input += f":{key}={value}"
        
        # Create SHA-256 hash of the input
        cache_key = hashlib.sha256(cache_input.encode()).hexdigest()[:16]
        self.logger.debug(f"Generated cache key {cache_key} for operation: {operation}")
        return cache_key
    
    def _get_cached_response(self, cache_key: str) -> Dict:
        """Retrieve cached response if it exists and is valid"""
        cache_file = self.cache_dir / f"{cache_key}.json"
        
        if not cache_file.exists():
            self.logger.debug(f"No cache file found for key: {cache_key}")
            return {}
        
        try:
            with open(cache_file, 'r', encoding='utf-8') as f:
                cached_data = json.load(f)
            
            # Check if cache is expired
            cache_time = datetime.fromisoformat(cached_data.get('timestamp', ''))
            expiration_days = self.config.get_cache_expiration_days()
            if (datetime.now() - cache_time).days < expiration_days:
                self.logger.info(f"Using cached response for {cache_key[:8]}...")
                self.logger.debug(f"Cache hit - file: {cache_file}, age: {(datetime.now() - cache_time).days} days")
                return cached_data.get('response', {})
            else:
                # Cache expired, remove it
                cache_file.unlink()
                self.logger.info(f"Expired cache removed for {cache_key[:8]} (age: {(datetime.now() - cache_time).days} days)")
        except (json.JSONDecodeError, KeyError, ValueError, OSError) as e:
            # Invalid or corrupted cache file, remove it
            try:
                cache_file.unlink()
                self.logger.warning(f"Corrupted cache removed for {cache_key[:8]}: {e}")
            except OSError:
                self.logger.error(f"Failed to remove corrupted cache file: {cache_file}")
        
        return {}
    
    def _save_cached_response(self, cache_key: str, response: Dict) -> None:
        """Save response to cache"""
        try:
            cache_file = self.cache_dir / f"{cache_key}.json"
            cache_data = {
                'timestamp': datetime.now().isoformat(),
                'response': response
            }
            
            with open(cache_file, 'w', encoding='utf-8') as f:
                json.dump(cache_data, f, indent=2)
            
            self.logger.info(f"Cached response for {cache_key[:8]}...")
            self.logger.debug(f"Cache saved to: {cache_file}")
        except Exception as e:
            self.logger.error(f"Could not save cache for {cache_key[:8]}: {e}")
    
    def anonymize_resume(self, resume_content: str) -> str:
        """Remove or anonymize PII from resume content"""
        self.logger.debug("Starting PII anonymization process")
        content = resume_content
        
        # Check if PII removal is enabled
        if not self.config.get('resume_processing.pii_removal.enabled', True):
            self.logger.info("PII removal disabled in configuration")
            return content
        
        # Track what we're removing for debugging
        pii_removed = []
        
        # 1. Email addresses
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails_found = re.findall(email_pattern, content)
        if emails_found:
            content = re.sub(email_pattern, '[EMAIL_REDACTED]', content)
            pii_removed.append(f"{len(emails_found)} email(s)")
            self.logger.debug(f"Found and redacted {len(emails_found)} email addresses")
        
        # 2. Phone numbers (various formats)
        phone_patterns = [
            r'\b\(?(\d{3})\)?[-.\s]?(\d{3})[-.\s]?(\d{4})\b',  # (555) 123-4567, 555-123-4567, 555.123.4567
            r'\b(\d{3})[-.\s](\d{3})[-.\s](\d{4})\b',          # 555 123 4567
            r'\+1[-.\s]?\(?(\d{3})\)?[-.\s]?(\d{3})[-.\s]?(\d{4})\b'  # +1 (555) 123-4567
        ]
        
        for i, pattern in enumerate(phone_patterns):
            phones_found = re.findall(pattern, content)
            if phones_found:
                content = re.sub(pattern, '[PHONE_REDACTED]', content)
                pii_removed.append(f"{len(phones_found)} phone number(s)")
                self.logger.debug(f"Found and redacted {len(phones_found)} phone numbers with pattern {i+1}")
        
        # 3. Physical addresses (basic patterns)
        # Remove lines that look like addresses (number + street + city/state/zip patterns)
        address_patterns = [
            r'\b\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Lane|Ln|Drive|Dr|Court|Ct|Place|Pl)\b.*',
            r'\b[A-Za-z\s]+,\s*[A-Z]{2}\s+\d{5}(?:-\d{4})?\b'  # City, ST 12345 or City, ST 12345-6789
        ]
        
        for i, pattern in enumerate(address_patterns):
            addresses_found = re.findall(pattern, content)
            if addresses_found:
                content = re.sub(pattern, '[ADDRESS_REDACTED]', content)
                pii_removed.append(f"{len(addresses_found)} address(es)")
                self.logger.debug(f"Found and redacted {len(addresses_found)} addresses with pattern {i+1}")
        
        # 4. Personal websites/portfolios
        url_pattern = r'https?://(?:www\.)?[A-Za-z0-9.-]+\.[A-Za-z]{2,}(?:/[^\s]*)?'
        urls_found = re.findall(url_pattern, content)
        if urls_found:
            # Only redact personal domains if preserve_professional_urls is enabled
            if self.config.get('resume_processing.pii_removal.preserve_professional_urls', True):
                professional_domains = self.config.get_professional_domains()
                personal_urls = []
                for url in urls_found:
                    is_professional = any(domain in url.lower() for domain in professional_domains)
                    if not is_professional:
                        content = content.replace(url, '[WEBSITE_REDACTED]')
                        personal_urls.append(url)
                if personal_urls:
                    pii_removed.append(f"{len(personal_urls)} personal website(s)")
                    self.logger.debug(f"Redacted {len(personal_urls)} personal URLs, preserved {len(urls_found) - len(personal_urls)} professional URLs")
            else:
                # Redact all URLs
                for url in urls_found:
                    content = content.replace(url, '[WEBSITE_REDACTED]')
                pii_removed.append(f"{len(urls_found)} website(s)")
                self.logger.debug(f"Redacted all {len(urls_found)} URLs")
        
        # 5. Names (more complex - try to identify the name at the top of resume)
        lines = content.split('\n')
        if lines:
            first_line = lines[0].strip()
            # If first line looks like a name (2-3 words, title case, no numbers)
            if (len(first_line.split()) in [2, 3] and 
                first_line.replace(' ', '').isalpha() and 
                first_line.istitle() and
                len(first_line) < 50):
                lines[0] = '[NAME_REDACTED]'
                content = '\n'.join(lines)
                pii_removed.append("name")
                self.logger.debug(f"Redacted candidate name from first line: {first_line}")
        
        if pii_removed:
            self.logger.info(f"PII removed: {', '.join(pii_removed)}")
        else:
            self.logger.info("No PII detected in resume")
        
        self.logger.debug("PII anonymization process completed")
        return content
    
    def read_resume_file(self, file_path: str) -> str:
        """Read resume content from various file formats"""
        self.logger.info(f"Reading resume file: {file_path}")
        file_extension = file_path.lower().split('.')[-1]
        
        try:
            if file_extension == 'txt':
                self.logger.debug("Processing TXT file")
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()
                    self.logger.debug(f"TXT file read successfully - {len(content)} characters")
                    return content
            
            elif file_extension == 'pdf':
                self.logger.debug("Processing PDF file")
                text = ""
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    page_count = len(pdf_reader.pages)
                    self.logger.debug(f"PDF has {page_count} pages")
                    for i, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        text += page_text
                        self.logger.debug(f"Extracted {len(page_text)} characters from page {i+1}")
                self.logger.debug(f"PDF file processed successfully - {len(text)} total characters")
                return text
            
            elif file_extension in ['docx', 'doc']:
                self.logger.debug(f"Processing {file_extension.upper()} file")
                doc = Document(file_path)
                text = []
                paragraph_count = len(doc.paragraphs)
                self.logger.debug(f"Document has {paragraph_count} paragraphs")
                for paragraph in doc.paragraphs:
                    text.append(paragraph.text)
                content = '\n'.join(text)
                self.logger.debug(f"DOCX file processed successfully - {len(content)} characters")
                return content
            
            else:
                error_msg = f"Unsupported file format: {file_extension}"
                self.logger.error(error_msg)
                raise ValueError(error_msg)
                
        except Exception as e:
            self.logger.error(f"Error reading resume file {file_path}: {str(e)}")
            raise
    
    def extract_keywords(self, resume_content: str) -> Dict:
        """Extract relevant keywords and information from resume using OpenAI"""
        self.logger.info("Starting keyword extraction from resume")
        
        # Anonymize the resume content before sending to API
        anonymized_content = self.anonymize_resume(resume_content)
        self.logger.debug(f"Resume anonymized - content length: {len(anonymized_content)} characters")
        
        # Check cache first
        cache_key = self._generate_cache_key(anonymized_content, "extract_keywords")
        cached_response = self._get_cached_response(cache_key)
        if cached_response:
            self.logger.info("Using cached keyword extraction results")
            return cached_response
        
        prompt = f"""
        Analyze the following resume and extract key information that would be useful for job searching:

        Resume Content:
        {anonymized_content}

        Please extract and categorize the following information in JSON format:
        1. Technical skills (programming languages, frameworks, tools, databases, cloud platforms)
        2. Job titles/roles the person has held or would be suitable for
        3. Years of experience (estimate if not explicitly stated)
        4. Industry/domain expertise
        5. Key achievements or specializations
        6. Location preferences (if mentioned, otherwise use "Not specified")

        Format your response as a JSON object with the following structure:
        {{
            "technical_skills": ["skill1", "skill2", ...],
            "job_titles": ["title1", "title2", ...],
            "years_of_experience": "X years",
            "industries": ["industry1", "industry2", ...],
            "specializations": ["spec1", "spec2", ...],
            "location": "location if mentioned or Not specified"
        }}
        """
        
        try:
            self.logger.debug("Sending keyword extraction request to OpenAI API")
            response = self.client.chat.completions.create(
                model=self.config.get_openai_model(),
                messages=[
                    {"role": "system", "content": "You are an expert HR professional and career counselor. Extract key information from resumes accurately and format it as requested. Note that some PII has been redacted for privacy."},
                    {"role": "user", "content": prompt}
                ],
                temperature=self.config.get_openai_temperature()
            )
            
            self.logger.debug("Received response from OpenAI API")
            content = response.choices[0].message.content
            
            # Try to parse JSON from the response
            try:
                result = json.loads(content)
                self.logger.info("Successfully parsed keyword extraction JSON response")
                self.logger.debug(f"Extracted keywords: {list(result.keys())}")
            except json.JSONDecodeError:
                self.logger.warning("Initial JSON parsing failed, trying to extract from markdown code blocks")
                # If JSON parsing fails, try to extract JSON from markdown code blocks
                import re
                json_match = re.search(r'```json\s*(.*?)\s*```', content, re.DOTALL)
                if json_match:
                    result = json.loads(json_match.group(1))
                    self.logger.info("Successfully extracted JSON from markdown code blocks")
                else:
                    # If still fails, return the raw content for debugging
                    self.logger.error("Failed to parse JSON response, returning raw content")
                    result = {"raw_response": content}
            
            # Cache the result
            self._save_cached_response(cache_key, result)
            self.logger.info("Keyword extraction completed successfully")
            return result
                    
        except Exception as e:
            self.logger.error(f"Error extracting keywords: {e}")
            return {}
    
    def generate_search_terms(self, keywords_data: Dict, target_location: str = None, desired_position: str = None) -> Dict:
        """Generate optimized search terms for job boards based on extracted keywords"""
        self.logger.info("Generating optimized search terms")
        self.logger.debug(f"Input: target_location={target_location}, desired_position={desired_position}")
        
        # Check cache first
        cache_key = self._generate_cache_key(
            json.dumps(keywords_data, sort_keys=True), 
            "generate_search_terms",
            location=target_location or "",
            position=desired_position or ""
        )
        cached_response = self._get_cached_response(cache_key)
        if cached_response:
            self.logger.info("Using cached search terms generation results")
            return cached_response
        
        desired_position_context = f"\nDesired Position: {desired_position}" if desired_position else ""
        
        prompt = f"""
        Based on the following extracted resume information, generate optimized search terms for job board scraping:

        Resume Data:
        {json.dumps(keywords_data, indent=2)}

        Target Location: {target_location or "Not specified"}{desired_position_context}

        Please generate the following search parameters in JSON format:
        1. Primary search terms (2-3 most relevant job titles/roles{' - prioritize the desired position if provided' if desired_position else ''})
        2. Secondary search terms (broader terms that might capture relevant jobs)
        3. Skills-based search terms (combinations of key technical skills)
        4. Suggested location (use target_location if provided, otherwise extract from resume)
        5. Experience level filter suggestions
        6. Google search optimization string (for sites that support it)

        {f'IMPORTANT: The user specifically wants to target "{desired_position}" roles. Please prioritize this position in your search terms while still considering the resume skills and experience.' if desired_position else ''}

        Format your response as a JSON object:
        {{
            "primary_search_terms": ["term1", "term2", "term3"],
            "secondary_search_terms": ["term1", "term2", "term3"],
            "skills_based_terms": ["skill combo 1", "skill combo 2"],
            "location": "suggested location",
            "experience_level": "junior/mid/senior",
            "google_search_string": "optimized search string for Google job search"
        }}
        """
        
        try:
            self.logger.debug("Sending search terms generation request to OpenAI API")
            response = self.client.chat.completions.create(
                model=self.config.get_openai_model(),
                messages=[
                    {"role": "system", "content": "You are an expert recruiter who understands how to optimize job search queries for maximum relevant results. When a desired position is specified, prioritize it while leveraging the candidate's existing skills."},
                    {"role": "user", "content": prompt}
                ],
                temperature=self.config.get_openai_temperature()
            )
            
            self.logger.debug("Received response from OpenAI API")
            content = response.choices[0].message.content
            try:
                result = json.loads(content)
                self.logger.info("Successfully parsed search terms generation JSON response")
            except json.JSONDecodeError:
                self.logger.warning("Initial JSON parsing failed, trying to extract from markdown code blocks")
                import re
                json_match = re.search(r'```json\s*(.*?)\s*```', content, re.DOTALL)
                if json_match:
                    result = json.loads(json_match.group(1))
                    self.logger.info("Successfully extracted JSON from markdown code blocks")
                else:
                    self.logger.error("Failed to parse JSON response, returning raw content")
                    result = {"raw_response": content}
            
            # Cache the result
            self._save_cached_response(cache_key, result)
            self.logger.info("Search terms generation completed successfully")
            return result
                    
        except Exception as e:
            self.logger.error(f"Error generating search terms: {e}")
            return {}
    
    def clear_cache(self) -> None:
        """Clear all cached responses"""
        self.logger.info("Starting cache cleanup")
        try:
            import shutil
            if self.cache_dir.exists():
                # Remove all .json files in cache directory
                cache_files = list(self.cache_dir.glob("*.json"))
                deleted_count = 0
                for cache_file in cache_files:
                    try:
                        cache_file.unlink()
                        deleted_count += 1
                        self.logger.debug(f"Deleted cache file: {cache_file}")
                    except Exception as e:
                        self.logger.error(f"Could not delete cache file {cache_file}: {e}")
                
                self.logger.info(f"Cache cleared - {deleted_count} files deleted")
            else:
                self.logger.warning("Cache directory doesn't exist")
        except Exception as e:
            self.logger.error(f"Error clearing cache: {e}")
    
    def get_cache_info(self) -> Dict:
        """Get information about cached responses"""
        self.logger.debug("Gathering cache information")
        if not self.cache_dir.exists():
            self.logger.debug("Cache directory doesn't exist")
            return {
                "cache_files_count": 0,
                "total_size_mb": 0,
                "cache_directory": str(self.cache_dir),
                "cache_files": []
            }
        
        try:
            cache_files = list(self.cache_dir.glob("*.json"))
            total_size = sum(f.stat().st_size for f in cache_files if f.is_file())
            
            self.logger.debug(f"Found {len(cache_files)} cache files, total size: {total_size} bytes")
            
            # Get detailed info about each cache file
            cache_file_details = []
            expiration_days = self.config.get_cache_expiration_days()
            
            for cache_file in cache_files:
                try:
                    # Read the cache file to get timestamp and check if expired
                    with open(cache_file, 'r', encoding='utf-8') as f:
                        cache_data = json.load(f)
                    
                    created_time = datetime.fromisoformat(cache_data.get('timestamp', ''))
                    is_expired = (datetime.now() - created_time).days > expiration_days
                    
                    cache_file_details.append({
                        'key': cache_file.stem,  # filename without extension
                        'created': created_time.strftime('%Y-%m-%d %H:%M:%S'),
                        'size_kb': round(cache_file.stat().st_size / 1024, 2),
                        'is_expired': is_expired
                    })
                except Exception as e:
                    # If we can't read a cache file, just show basic info
                    self.logger.warning(f"Error reading cache file {cache_file}: {e}")
                    cache_file_details.append({
                        'key': cache_file.stem,
                        'created': 'Unknown',
                        'size_kb': round(cache_file.stat().st_size / 1024, 2) if cache_file.exists() else 0,
                        'is_expired': False,
                        'error': str(e)
                    })
            
            cache_info = {
                "cache_files_count": len(cache_files),
                "total_size_mb": round(total_size / (1024 * 1024), 2),
                "cache_directory": str(self.cache_dir),
                "cache_files": cache_file_details
            }
            
            self.logger.debug(f"Cache info gathered: {cache_info['cache_files_count']} files, {cache_info['total_size_mb']} MB")
            return cache_info
            
        except Exception as e:
            self.logger.error(f"Error gathering cache info: {e}")
            return {
                "cache_files_count": 0,
                "total_size_mb": 0,
                "cache_directory": str(self.cache_dir),
                "cache_files": [],
                "error": str(e)
            }
    
    def process_resume(self, resume_file_path: str, target_location: str = None, desired_position: str = None) -> Dict:
        """Complete pipeline: read resume, extract keywords, generate search terms"""
        
        self.logger.info(f"Starting complete resume processing pipeline for: {resume_file_path}")
        if desired_position:
            self.logger.info(f"Targeting position: {desired_position}")
        if target_location:
            self.logger.info(f"Target location: {target_location}")
        
        # Show cache info
        cache_info = self.get_cache_info()
        if cache_info["cache_files_count"] > 0:
            self.logger.info(f"Cache status: {cache_info['cache_files_count']} files, {cache_info['total_size_mb']} MB")
        
        try:
            # Step 1: Read resume content
            self.logger.info("Step 1: Reading resume content")
            resume_content = self.read_resume_file(resume_file_path)
            self.logger.info(f"Resume content loaded ({len(resume_content)} characters)")
            
            # Step 2: Extract keywords
            self.logger.info("Step 2: Extracting keywords from resume")
            keywords_data = self.extract_keywords(resume_content)
            
            # Step 3: Generate search terms
            self.logger.info("Step 3: Generating optimized search terms")
            search_terms = self.generate_search_terms(keywords_data, target_location, desired_position)
            
            result = {
                "keywords": keywords_data,
                "search_terms": search_terms,
                "resume_length": len(resume_content),
                "desired_position": desired_position
            }
            
            self.logger.info("Resume processing pipeline completed successfully")
            self.logger.debug(f"Result keys: {list(result.keys())}")
            return result
            
        except Exception as e:
            self.logger.error(f"Error in resume processing pipeline: {str(e)}")
            raise
    
    def analyze_and_rank_jobs(self, jobs_data: List[Dict], resume_keywords: Dict, max_jobs: int = None) -> List[Dict]:
        """
        Analyze jobs for salary information and rank by similarity to resume.
        
        Args:
            jobs_data: List of job dictionaries from job scraping
            resume_keywords: Keywords extracted from resume
            max_jobs: Maximum number of jobs to analyze (cost control)
            
        Returns:
            List of job dictionaries with added analysis data
        """
        start_time = time.time()
        self.logger.info(f"Starting job analysis for {len(jobs_data)} jobs")
        
        if not self.config.get_job_analysis_enabled():
            self.logger.info("Job analysis disabled in configuration")
            return jobs_data
        
        # Limit number of jobs to analyze for cost control
        analysis_limit = max_jobs or self.config.get_max_jobs_to_analyze()
        jobs_to_analyze = jobs_data[:analysis_limit]
        
        self.logger.info(f"Analyzing {len(jobs_to_analyze)} jobs (limited from {len(jobs_data)} for cost control)")
        
        analyzed_jobs = []
        batch_size = self.config.get_job_analysis_batch_size()
        total_batches = (len(jobs_to_analyze) + batch_size - 1) // batch_size
        
        self.logger.debug(f"Processing jobs in {total_batches} batches of size {batch_size}")
        
        # Check if parallel processing is enabled
        if self.config.get_job_analysis_parallel_enabled() and total_batches > 1:
            self.logger.info(f"Using parallel processing for {total_batches} batches")
            analyzed_jobs = self._process_batches_parallel(jobs_to_analyze, batch_size, resume_keywords)
        else:
            self.logger.info("Using sequential processing")
            analyzed_jobs = self._process_batches_sequential(jobs_to_analyze, batch_size, resume_keywords)
        
        # Add unanalyzed jobs back to the list with default scores
        remaining_jobs = jobs_data[analysis_limit:]
        if remaining_jobs:
            self.logger.info(f"Adding {len(remaining_jobs)} unanalyzed jobs with default scores")
            for job in remaining_jobs:
                job['analyzed'] = False
                job['similarity_score'] = 0.0
                job['salary_min_extracted'] = None
                job['salary_max_extracted'] = None
                job['salary_confidence'] = 0.0
                analyzed_jobs.append(job)
        
        # Sort by similarity score (highest first) if similarity ranking is enabled
        if self.config.get_similarity_ranking_enabled():
            original_order = [job.get('similarity_score', 0) for job in analyzed_jobs[:5]]  # Log first 5 for comparison
            analyzed_jobs.sort(key=lambda x: x.get('similarity_score', 0), reverse=True)
            new_order = [job.get('similarity_score', 0) for job in analyzed_jobs[:5]]
            self.logger.info("Jobs ranked by similarity score")
            self.logger.debug(f"Top 5 similarity scores before sorting: {original_order}")
            self.logger.debug(f"Top 5 similarity scores after sorting: {new_order}")
        
        elapsed_time = time.time() - start_time
        self.logger.info(f"Job analysis completed in {elapsed_time:.2f} seconds - {len(analyzed_jobs)} total jobs processed")
        return analyzed_jobs
    
    def _process_batches_parallel(self, jobs_to_analyze: List[Dict], batch_size: int, resume_keywords: Dict) -> List[Dict]:
        """Process job batches in parallel"""
        self.logger.debug(f"Processing {len(jobs_to_analyze)} jobs in parallel")
        
        analyzed_jobs = []
        batch_data = []
        
        # Prepare batch data with indices to maintain order
        for i in range(0, len(jobs_to_analyze), batch_size):
            batch = jobs_to_analyze[i:i + batch_size]
            batch_data.append((i // batch_size, batch))
        
        max_workers = self.config.get_job_analysis_parallel_workers()
        delay_between_requests = self.config.get_job_analysis_request_delay()
        
        self.logger.info(f"Processing {len(batch_data)} batches with max {max_workers} parallel workers")
        
        # Create a ThreadPoolExecutor to process job batches in parallel
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Submit all batch jobs with a delay between submissions to be API-friendly
            futures_to_batch = {}
            for batch_idx, batch in batch_data:
                future = executor.submit(self._analyze_job_batch_with_delay, batch, resume_keywords, delay_between_requests)
                futures_to_batch[future] = batch_idx
                
                # Small delay between submissions to be respectful
                if batch_idx < len(batch_data) - 1:
                    time.sleep(0.1)
            
            # Collect results in order
            batch_results = [None] * len(batch_data)
            for future in as_completed(futures_to_batch):
                batch_idx = futures_to_batch[future]
                try:
                    batch_result = future.result()
                    batch_results[batch_idx] = batch_result
                    self.logger.debug(f"Completed batch {batch_idx + 1}/{len(batch_data)}")
                except Exception as e:
                    self.logger.error(f"Error in parallel batch {batch_idx + 1}: {str(e)}")
                    # Use default analysis for failed batch
                    batch_results[batch_idx] = self._create_default_analysis(batch_data[batch_idx][1])
            
            # Combine results from all batches in order
            for batch_result in batch_results:
                if batch_result:
                    analyzed_jobs.extend(batch_result)
        
        self.logger.info(f"Parallel processing completed - processed {len(analyzed_jobs)} jobs")
        return analyzed_jobs
    
    def _analyze_job_batch_with_delay(self, jobs_batch: List[Dict], resume_keywords: Dict, delay: float) -> List[Dict]:
        """Analyze a batch of jobs with an optional delay for rate limiting"""
        if delay > 0:
            time.sleep(delay)
        return self._analyze_job_batch(jobs_batch, resume_keywords)
    
    def _process_batches_sequential(self, jobs_to_analyze: List[Dict], batch_size: int, resume_keywords: Dict) -> List[Dict]:
        """Process job batches sequentially"""
        self.logger.debug(f"Processing {len(jobs_to_analyze)} jobs sequentially")
        
        analyzed_jobs = []
        for i in range(0, len(jobs_to_analyze), batch_size):
            batch = jobs_to_analyze[i:i + batch_size]
            analyzed_jobs.extend(self._analyze_job_batch(batch, resume_keywords))
        
        self.logger.debug(f"Processed {len(analyzed_jobs)} jobs sequentially")
        return analyzed_jobs
    
    def _analyze_job_batch(self, jobs_batch: List[Dict], resume_keywords: Dict) -> List[Dict]:
        """Analyze a batch of jobs for salary and similarity"""
        self.logger.debug(f"Analyzing batch of {len(jobs_batch)} jobs")
        
        # Prepare resume summary for comparison
        resume_summary = self._create_resume_summary(resume_keywords)
        self.logger.debug(f"Created resume summary: {resume_summary[:100]}...")
        
        # Prepare job descriptions for analysis
        jobs_for_analysis = []
        for i, job in enumerate(jobs_batch):
            # Safely handle description field that might be float/NaN
            description = job.get('description', 'N/A')
            if not isinstance(description, str):
                # Convert non-string values (like float/NaN) to string
                if pd.isna(description):
                    description = 'N/A'
                else:
                    description = str(description)
            
            job_text = f"""
            Title: {job.get('title', 'N/A')}
            Company: {job.get('company', 'N/A')}
            Location: {job.get('location', 'N/A')}
            Description: {description[:1000]}  # Limit description length
            """
            jobs_for_analysis.append({
                'index': i,
                'job_text': job_text.strip()
            })
        
        self.logger.debug(f"Prepared {len(jobs_for_analysis)} jobs for analysis")
        
        # Check cache first
        cache_key = self._generate_cache_key(
            json.dumps([j['job_text'] for j in jobs_for_analysis], sort_keys=True),
            "analyze_jobs_batch",
            resume_summary=resume_summary
        )
        cached_response = self._get_cached_response(cache_key)
        if cached_response:
            self.logger.info("Using cached job analysis results")
            return self._apply_analysis_to_jobs(jobs_batch, cached_response)
        
        # Create analysis prompt
        prompt = self._create_job_analysis_prompt(jobs_for_analysis, resume_summary)
        
        try:
            self.logger.debug("Sending job analysis request to OpenAI API")
            response = self.client.chat.completions.create(
                model=self.config.get_job_analysis_model(),
                messages=[
                    {"role": "system", "content": "You are an expert HR analyst and career counselor. Analyze job postings for salary information and similarity to candidate profiles. Be accurate and conservative in your assessments."},
                    {"role": "user", "content": prompt}
                ],
                temperature=self.config.get_openai_temperature()
            )
            
            self.logger.debug("Received job analysis response from OpenAI API")
            content = response.choices[0].message.content
            try:
                analysis_results = json.loads(content)
                self.logger.debug("Successfully parsed job analysis JSON response")
            except json.JSONDecodeError:
                self.logger.warning("Initial JSON parsing failed, trying to extract from markdown code blocks")
                import re
                json_match = re.search(r'```json\s*(.*?)\s*```', content, re.DOTALL)
                if json_match:
                    analysis_results = json.loads(json_match.group(1))
                    self.logger.info("Successfully extracted JSON from markdown code blocks")
                else:
                    self.logger.error(f"Failed to parse job analysis response: {content}")
                    return self._create_default_analysis(jobs_batch)
            
            # Cache the result
            if self.config.get('job_analysis.cache_analysis_results', True):
                self._save_cached_response(cache_key, analysis_results)
            
            self.logger.debug("Job batch analysis completed successfully")
            return self._apply_analysis_to_jobs(jobs_batch, analysis_results)
            
        except Exception as e:
            self.logger.error(f"Error analyzing job batch: {e}")
            return self._create_default_analysis(jobs_batch)
    
    def _create_resume_summary(self, resume_keywords: Dict) -> str:
        """Create a concise resume summary for job comparison"""
        self.logger.debug("Creating resume summary for job comparison")
        summary_parts = []
        
        if resume_keywords.get('technical_skills'):
            summary_parts.append(f"Technical Skills: {', '.join(resume_keywords['technical_skills'][:10])}")
        
        if resume_keywords.get('job_titles'):
            summary_parts.append(f"Experience: {', '.join(resume_keywords['job_titles'][:5])}")
        
        if resume_keywords.get('industries'):
            summary_parts.append(f"Industries: {', '.join(resume_keywords['industries'][:3])}")
        
        if resume_keywords.get('years_of_experience'):
            summary_parts.append(f"Years of Experience: {resume_keywords['years_of_experience']}")
        
        summary = ' | '.join(summary_parts)
        self.logger.debug(f"Created resume summary with {len(summary_parts)} components: {summary[:100]}...")
        return summary
    
    def _create_job_analysis_prompt(self, jobs_for_analysis: List[Dict], resume_summary: str) -> str:
        """Create the prompt for job analysis"""
        self.logger.debug(f"Creating job analysis prompt for {len(jobs_for_analysis)} jobs")
        
        jobs_text = ""
        for job_info in jobs_for_analysis:
            jobs_text += f"\n--- Job {job_info['index']} ---\n{job_info['job_text']}\n"
        
        analyze_salary = self.config.get_salary_analysis_enabled()
        analyze_similarity = self.config.get_similarity_ranking_enabled()
        
        self.logger.debug(f"Analysis settings - salary: {analyze_salary}, similarity: {analyze_similarity}")
        
        # Build the JSON template based on what analysis is enabled
        json_fields = []
        json_fields.append('"job_index": 0')
        
        if analyze_salary:
            json_fields.extend([
                '"salary_min": null',
                '"salary_max": null', 
                '"salary_confidence": 0.0'
            ])
        
        if analyze_similarity:
            json_fields.extend([
                '"similarity_score": 0.0',
                '"similarity_explanation": ""'
            ])
        
        json_fields.extend([
            '"key_matches": ["match1", "match2"]',
            '"missing_requirements": ["req1", "req2"]'
        ])
        
        json_template = "{\n                    " + ",\n                    ".join(json_fields) + "\n                }"
        
        # Build instructions based on enabled features
        instructions = []
        if analyze_salary:
            instructions.append("- Extract salary information if mentioned (annual salary in USD). Set confidence 0-1 based on how explicit the salary info is.")
        if analyze_similarity:
            instructions.append("- Rate similarity 0-10 based on skill match, experience level, and role alignment.")
        
        instructions.extend([
            "- Identify key skill/experience matches between candidate and job requirements",
            "- Note any significant missing requirements",
            "- Be conservative in salary extraction - only extract if clearly stated"
        ])
        
        if analyze_similarity:
            instructions.append("- Consider both technical skills and domain experience in similarity scoring")
        
        instructions_text = "\n        ".join(instructions)
        
        prompt = f"""
        Analyze the following job postings in relation to this candidate profile:
        
        CANDIDATE PROFILE:
        {resume_summary}
        
        JOB POSTINGS:{jobs_text}
        
        For each job, provide analysis in the following JSON format:
        {{
            "job_analyses": [
                {json_template}
            ]
        }}
        
        Instructions:
        {instructions_text}
        """
        
        self.logger.debug(f"Created job analysis prompt - {len(prompt)} characters")
        return prompt.strip()
    
    def _apply_analysis_to_jobs(self, jobs_batch: List[Dict], analysis_results: Dict) -> List[Dict]:
        """Apply analysis results to job dictionaries"""
        self.logger.debug(f"Applying analysis results to {len(jobs_batch)} jobs")
        analyzed_jobs = []
        
        job_analyses = analysis_results.get('job_analyses', [])
        self.logger.debug(f"Found {len(job_analyses)} job analyses in results")
        
        for i, job in enumerate(jobs_batch):
            job_copy = job.copy()
            job_copy['analyzed'] = True
            
            # Find corresponding analysis
            analysis = None
            for job_analysis in job_analyses:
                if job_analysis.get('job_index') == i:
                    analysis = job_analysis
                    break
            
            if analysis:
                self.logger.debug(f"Applying analysis to job {i}")
                # Apply salary analysis
                if self.config.get_salary_analysis_enabled():
                    job_copy['salary_min_extracted'] = analysis.get('salary_min')
                    job_copy['salary_max_extracted'] = analysis.get('salary_max')
                    job_copy['salary_confidence'] = analysis.get('salary_confidence', 0.0)
                
                # Apply similarity analysis with sanitization
                if self.config.get_similarity_ranking_enabled():
                    job_copy['similarity_score'] = analysis.get('similarity_score', 0.0)
                    # Sanitize explanation text for JSON safety
                    explanation = analysis.get('similarity_explanation', '')
                    if isinstance(explanation, str):
                        # Clean up problematic characters
                        explanation = explanation.replace('\n', ' ').replace('\r', ' ').replace('\t', ' ')
                        explanation = explanation.replace('"', "'").replace('\\', '/')
                        # Limit length
                        if len(explanation) > 500:
                            explanation = explanation[:500] + "..."
                    job_copy['similarity_explanation'] = explanation
                
                # Apply other analysis data with sanitization
                key_matches = analysis.get('key_matches', [])
                missing_requirements = analysis.get('missing_requirements', [])
                
                # Sanitize list items
                if isinstance(key_matches, list):
                    key_matches = [str(match).replace('"', "'").replace('\n', ' ')[:100] for match in key_matches if match]
                if isinstance(missing_requirements, list):
                    missing_requirements = [str(req).replace('"', "'").replace('\n', ' ')[:100] for req in missing_requirements if req]
                
                job_copy['key_matches'] = key_matches
                job_copy['missing_requirements'] = missing_requirements
            else:
                # Default values if no analysis found
                self.logger.warning(f"No analysis found for job {i}, using defaults")
                job_copy.update(self._get_default_job_analysis())
            
            analyzed_jobs.append(job_copy)
        
        self.logger.debug(f"Applied analysis to {len(analyzed_jobs)} jobs")
        return analyzed_jobs
    
    def _create_default_analysis(self, jobs_batch: List[Dict]) -> List[Dict]:
        """Create default analysis when API call fails"""
        self.logger.warning(f"Creating default analysis for {len(jobs_batch)} jobs due to API failure")
        analyzed_jobs = []
        for job in jobs_batch:
            job_copy = job.copy()
            job_copy['analyzed'] = False
            job_copy.update(self._get_default_job_analysis())
            analyzed_jobs.append(job_copy)
        return analyzed_jobs
    
    def _get_default_job_analysis(self) -> Dict:
        """Get default analysis values"""
        return {
            'similarity_score': 0.0,
            'similarity_explanation': 'Analysis not available',
            'salary_min_extracted': None,
            'salary_max_extracted': None,
            'salary_confidence': 0.0,
            'key_matches': [],
            'missing_requirements': []
        } 