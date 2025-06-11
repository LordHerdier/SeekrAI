import os
import logging
import PyPDF2
from docx import Document
from pathlib import Path


class FileReader:
    """Handles reading content from various file formats.
    
    This class provides functionality to read and extract text content from different
    file formats commonly used for resumes and job-related documents. Supported formats
    include plain text (.txt), PDF (.pdf), and Microsoft Word documents (.doc, .docx).
    
    The class implements robust error handling and logging to track file processing
    operations and handle encoding issues gracefully.
    
    Attributes:
        logger (logging.Logger): Logger instance for tracking file operations and errors.
    
    Example:
        >>> reader = FileReader()
        >>> content = reader.read_resume_file("resume.pdf")
        >>> print(len(content))
        1234
    """
    
    def __init__(self):
        """Initialize the FileReader with a configured logger.
        
        Sets up logging for the FileReader instance to track file processing
        operations, errors, and debug information.
        """
        self.logger = logging.getLogger(f"{__name__}.{self.__class__.__name__}")
    
    def read_resume_file(self, file_path: str) -> str:
        """Read resume content from various file formats.
        
        This is the main public method that determines the file type based on its
        extension and delegates to the appropriate private method for content extraction.
        Supports TXT, PDF, DOC, and DOCX file formats.
        
        Args:
            file_path (str): Path to the resume file to be read. Can be absolute or
                relative path. The file extension is used to determine the processing method.
        
        Returns:
            str: The extracted text content from the file. Empty or whitespace-only
                content will trigger appropriate error handling in the format-specific
                methods.
        
        Raises:
            ValueError: If the file format is not supported (extension not in 
                ['txt', 'pdf', 'doc', 'docx']) or if the file contains no extractable content.
            FileNotFoundError: If the specified file path does not exist.
            PermissionError: If the file cannot be accessed due to permission restrictions.
            Exception: For any other file reading errors, with the original exception
                logged and re-raised.
        
        Example:
            >>> reader = FileReader()
            >>> content = reader.read_resume_file("/path/to/resume.pdf")
            >>> print(content[:100])
            John Doe
            Software Engineer
            Experience: 5 years...
        
        Note:
            The method automatically detects file format based on extension and
            handles encoding issues for text files by attempting multiple encodings.
        """
        self.logger.info(f"Reading resume file: {file_path}")
        file_extension = file_path.lower().split('.')[-1]
        
        try:
            if file_extension == 'txt':
                self.logger.debug("Processing TXT file")
                return self._read_txt_file(file_path)
            elif file_extension == 'pdf':
                self.logger.debug("Processing PDF file")
                return self._read_pdf_file(file_path)
            elif file_extension in ['doc', 'docx']:
                self.logger.debug(f"Processing {file_extension.upper()} file")
                return self._read_docx_file(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            self.logger.error(f"Error reading file {file_path}: {str(e)}")
            raise
    
    def _read_txt_file(self, file_path: str) -> str:
        """Read content from a plain text file with encoding fallback.
        
        Attempts to read the text file using UTF-8 encoding first, and falls back
        to Latin-1 encoding if UTF-8 fails. This approach handles most common
        encoding scenarios for text files.
        
        Args:
            file_path (str): Path to the text file to be read.
        
        Returns:
            str: The complete text content of the file as a string.
        
        Raises:
            FileNotFoundError: If the specified file does not exist.
            PermissionError: If the file cannot be read due to permissions.
            OSError: For other file system related errors.
        
        Example:
            >>> reader = FileReader()
            >>> content = reader._read_txt_file("resume.txt")
            >>> print(f"Read {len(content)} characters")
            Read 1500 characters
        
        Note:
            This method first attempts UTF-8 encoding (most common for modern text files)
            and automatically falls back to Latin-1 if a UnicodeDecodeError occurs.
            The encoding used is logged for debugging purposes.
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                self.logger.debug(f"Successfully read TXT file, {len(content)} characters")
                return content
        except UnicodeDecodeError:
            # Try with different encoding if UTF-8 fails
            with open(file_path, 'r', encoding='latin-1') as file:
                content = file.read()
                self.logger.debug(f"Successfully read TXT file with latin-1 encoding, {len(content)} characters")
                return content
    
    def _read_pdf_file(self, file_path: str) -> str:
        """Read and extract text content from a PDF file.
        
        Uses PyPDF2 library to extract text from all pages of a PDF document.
        Handles multi-page documents and filters out empty pages. Provides detailed
        logging for each page processed and handles extraction errors gracefully.
        
        Args:
            file_path (str): Path to the PDF file to be read.
        
        Returns:
            str: Concatenated text content from all pages, with pages separated
                by newline characters. Only non-empty pages are included.
        
        Raises:
            ValueError: If the PDF contains no extractable text (may be image-based
                or corrupted), or if the file is not a valid PDF.
            FileNotFoundError: If the specified PDF file does not exist.
            PermissionError: If the PDF file cannot be opened due to permissions.
            PyPDF2.errors.PdfReadError: If the PDF file is corrupted or invalid.
        
        Example:
            >>> reader = FileReader()
            >>> content = reader._read_pdf_file("document.pdf")
            >>> pages = content.split('\n')
            >>> print(f"Extracted content from {len(pages)} pages")
            Extracted content from 3 pages
        
        Note:
            - Empty or image-only pages are skipped and logged as warnings
            - Text extraction failures on individual pages are logged but don't
              stop processing of remaining pages
            - The method validates that at least some text was extracted before returning
        """
        content = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            self.logger.debug(f"PDF has {num_pages} pages")
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():  # Only add non-empty pages
                        content.append(page_text)
                        self.logger.debug(f"Extracted text from page {page_num + 1}: {len(page_text)} characters")
                    else:
                        self.logger.warning(f"Page {page_num + 1} appears to be empty or image-only")
                except Exception as e:
                    self.logger.warning(f"Could not extract text from page {page_num + 1}: {e}")
        
        full_content = '\n'.join(content)
        self.logger.debug(f"Successfully read PDF file, total {len(full_content)} characters")
        
        if not full_content.strip():
            raise ValueError("PDF appears to contain no extractable text. It may be image-based or corrupted.")
        
        return full_content
    
    def _read_docx_file(self, file_path: str) -> str:
        """Read and extract text content from a Microsoft Word document.
        
        Uses the python-docx library to extract text from DOCX files by iterating
        through all paragraphs in the document. Filters out empty paragraphs and
        provides detailed logging about the extraction process.
        
        Args:
            file_path (str): Path to the DOCX file to be read. Also handles legacy
                DOC files if they can be opened by the python-docx library.
        
        Returns:
            str: Concatenated text content from all non-empty paragraphs,
                with paragraphs separated by newline characters.
        
        Raises:
            ValueError: If the DOCX file contains no text content, or if the file
                cannot be read due to format issues.
            FileNotFoundError: If the specified DOCX file does not exist.
            PermissionError: If the DOCX file cannot be opened due to permissions.
            docx.opc.exceptions.PackageNotFoundError: If the file is not a valid
                DOCX format or is corrupted.
        
        Example:
            >>> reader = FileReader()
            >>> content = reader._read_docx_file("resume.docx")
            >>> paragraphs = content.split('\n')
            >>> print(f"Document has {len(paragraphs)} paragraphs")
            Document has 15 paragraphs
        
        Note:
            - Only extracts text content; formatting, images, and tables are not processed
            - Empty paragraphs are automatically filtered out
            - The method validates that the document contains extractable text content
            - Legacy DOC files may work but DOCX format is recommended for best results
        """
        try:
            doc = Document(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    content.append(paragraph.text)
            
            full_content = '\n'.join(content)
            self.logger.debug(f"Successfully read DOCX file, {len(doc.paragraphs)} paragraphs, {len(full_content)} characters")
            
            if not full_content.strip():
                raise ValueError("DOCX file appears to contain no text content.")
            
            return full_content
            
        except Exception as e:
            self.logger.error(f"Error reading DOCX file: {e}")
            raise ValueError(f"Could not read DOCX file: {str(e)}") 